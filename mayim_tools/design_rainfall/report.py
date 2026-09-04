"""Report-ready output generation from a PointEstimate.

Two output formats:
    - CSV: flat duration x return-period table, for import into
      spreadsheets or other design workflows.
    - DOCX: a formatted Word table matching the layout engineers
      expect in a stormwater/drainage report, ready to paste or
      embed directly.
"""

from __future__ import annotations

import csv
from pathlib import Path

from .core import PointEstimate


def _duration_minutes(label: str) -> float:
    """Sort key: converts a duration label back to minutes so rows can
    be ordered chronologically (5 min < ... < 1 day < ... < 7 day)."""
    if label.endswith(" min"):
        return float(label[:-4])
    if label.endswith(" h"):
        return float(label[:-2]) * 60
    if label.endswith(" day"):
        return float(label[:-4]) * 1440
    return float("inf")


def _sorted_durations(results) -> list[str]:
    """Unique duration labels, chronologically ordered."""
    seen = {}
    for r in results:
        seen.setdefault(r.duration_label, _duration_minutes(r.duration_label))
    return sorted(seen, key=lambda d: seen[d])


def _sorted_return_periods(results) -> list[int]:
    return sorted(set(r.return_period for r in results))


def _write_stations_block(
    writer, stations_df, with_site_column: bool = False, site_label: str | None = None
) -> None:
    """Appends every field from a stations GeoDataFrame (as produced by
    DesignRainfallEngine.nearest_stations) - same data as the 'nearest
    rainfall stations' output layer, for reference/comparison only, not
    used in the grid-based calculation."""
    if stations_df is None or stations_df.empty:
        return
    cols = [c for c in stations_df.columns if c != "geometry"]
    writer.writerow((["Site"] if with_site_column else []) + cols)
    for _, row in stations_df.iterrows():
        writer.writerow(
            (([site_label] if with_site_column else [])) + [row[c] for c in cols]
        )


def write_csv(estimate: PointEstimate, output_path: str | Path, stations=None) -> None:
    """Wide format: one row per duration, one Depth/Lower/Upper column
    triplet per return period. If a stations GeoDataFrame is supplied
    (nearest daily rainfall stations, reference only), it's appended at
    the bottom of the file."""
    durations = _sorted_durations(estimate.results)
    return_periods = _sorted_return_periods(estimate.results)
    lookup = {(r.duration_label, r.return_period): r for r in estimate.results}

    with open(output_path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(
            ["Latitude", "Longitude", "MAP (mm)", "Altitude (m)", "Cluster"]
        )
        writer.writerow(
            [
                estimate.latitude,
                estimate.longitude,
                estimate.map_mm,
                estimate.altitude_m,
                estimate.cluster,
            ]
        )
        writer.writerow([])

        header = ["Duration"]
        for rt in return_periods:
            header += [f"{rt}yr Depth (mm)", f"{rt}yr Lower (mm)", f"{rt}yr Upper (mm)"]
        writer.writerow(header)

        for dur in durations:
            row = [dur]
            for rt in return_periods:
                r = lookup.get((dur, rt))
                if r:
                    row += [
                        r.depth,
                        r.lower if r.lower is not None else "",
                        r.upper if r.upper is not None else "",
                    ]
                else:
                    row += ["", "", ""]
            writer.writerow(row)

        if stations is not None and not stations.empty:
            writer.writerow([])
            writer.writerow(
                [
                    "Nearest Rainfall Stations (reference only - not used in the grid-based calculation)"
                ]
            )
            _write_stations_block(writer, stations)


def write_docx(
    estimate: PointEstimate,
    output_path: str | Path,
    site_name: str = "Site of interest",
) -> None:
    """Requires python-docx. Kept as a soft import so core.py / CSV
    export work without it installed."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Design Rainfall Estimate", level=1)
    doc.add_paragraph(
        f"Location: {site_name}  |  Lat: {estimate.latitude:.4f}, "
        f"Lon: {estimate.longitude:.4f}  |  MAP: {estimate.map_mm:.0f} mm  |  "
        f"Altitude: {estimate.altitude_m:.0f} m"
    )
    doc.add_paragraph(
        "Estimated using the regional L-moment design rainfall methodology "
        "(Smithers & Schulze), rebuilt from the original rainfall2/rainfall3 dataset."
    )

    # Pivot: rows = duration, columns = return period
    durations = sorted(
        set(r.duration_label for r in estimate.results),
        key=lambda d: [
            r.return_period for r in estimate.results if r.duration_label == d
        ][0],
    )
    return_periods = sorted(set(r.return_period for r in estimate.results))

    table = doc.add_table(rows=len(durations) + 1, cols=len(return_periods) + 1)
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    hdr[0].text = "Duration"
    for j, rt in enumerate(return_periods, start=1):
        hdr[j].text = f"{rt} yr"

    lookup = {(r.duration_label, r.return_period): r.depth for r in estimate.results}
    for i, dur in enumerate(durations, start=1):
        row = table.rows[i].cells
        row[0].text = dur
        for j, rt in enumerate(return_periods, start=1):
            depth = lookup.get((dur, rt))
            row[j].text = f"{depth:.1f}" if depth is not None else "-"

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.save(str(output_path))


# ----------------------------------------------------------------------
# Multi-site variants (for vector point layer input)
# ----------------------------------------------------------------------


def write_csv_multi(
    estimates: dict[str, PointEstimate],
    output_path: str | Path,
    stations_by_site: dict | None = None,
) -> None:
    """One combined CSV: a Site column identifies which point each row
    belongs to. Same wide layout as write_csv - one row per site+duration,
    one Depth/Lower/Upper column triplet per return period. The return
    period column set is the union across all sites, so every site's rows
    line up under the same headers even if individual sites happen to be
    missing a particular return period. If stations_by_site is supplied
    (dict[site_label, GeoDataFrame]), the nearest stations for every site
    are appended at the bottom of the file."""
    all_results = [r for est in estimates.values() for r in est.results]
    return_periods = _sorted_return_periods(all_results)

    with open(output_path, "w", newline="") as f:
        writer = csv.writer(f)
        header = [
            "Site",
            "Latitude",
            "Longitude",
            "MAP (mm)",
            "Altitude (m)",
            "Cluster",
            "Duration",
        ]
        for rt in return_periods:
            header += [f"{rt}yr Depth (mm)", f"{rt}yr Lower (mm)", f"{rt}yr Upper (mm)"]
        writer.writerow(header)

        for site, est in estimates.items():
            durations = _sorted_durations(est.results)
            lookup = {(r.duration_label, r.return_period): r for r in est.results}
            for dur in durations:
                row = [
                    site,
                    est.latitude,
                    est.longitude,
                    est.map_mm,
                    est.altitude_m,
                    est.cluster,
                    dur,
                ]
                for rt in return_periods:
                    r = lookup.get((dur, rt))
                    if r:
                        row += [
                            r.depth,
                            r.lower if r.lower is not None else "",
                            r.upper if r.upper is not None else "",
                        ]
                    else:
                        row += ["", "", ""]
                writer.writerow(row)

        if stations_by_site and any(not df.empty for df in stations_by_site.values()):
            writer.writerow([])
            writer.writerow(
                [
                    "Nearest Rainfall Stations (reference only - not used in the grid-based calculation)"
                ]
            )
            first = True
            for site, df in stations_by_site.items():
                if df is None or df.empty:
                    continue
                if first:
                    cols = [c for c in df.columns if c != "geometry"]
                    writer.writerow(["Site"] + cols)
                    first = False
                for _, row in df.iterrows():
                    writer.writerow([site] + [row[c] for c in cols])


def write_docx_multi(
    estimates: dict[str, PointEstimate], output_path: str | Path
) -> None:
    """One Word document, one heading + table per site."""
    from docx import Document

    doc = Document()
    doc.add_heading("Design Rainfall Estimates", level=1)
    doc.add_paragraph(
        "Estimated using the regional L-moment design rainfall methodology "
        "(Smithers & Schulze), rebuilt from the original rainfall2/rainfall3 dataset."
    )
    for site, est in estimates.items():
        _append_site_section(doc, site, est)
    doc.save(str(output_path))


def _append_site_section(doc, site_name: str, estimate: PointEstimate) -> None:
    from docx.shared import Pt

    doc.add_heading(site_name, level=2)
    doc.add_paragraph(
        f"Lat: {estimate.latitude:.4f}, Lon: {estimate.longitude:.4f}  |  "
        f"MAP: {estimate.map_mm:.0f} mm  |  Altitude: {estimate.altitude_m:.0f} m"
    )

    durations = sorted(
        set(r.duration_label for r in estimate.results),
        key=lambda d: [
            r.return_period for r in estimate.results if r.duration_label == d
        ][0],
    )
    return_periods = sorted(set(r.return_period for r in estimate.results))

    table = doc.add_table(rows=len(durations) + 1, cols=len(return_periods) + 1)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Duration"
    for j, rt in enumerate(return_periods, start=1):
        hdr[j].text = f"{rt} yr"

    lookup = {(r.duration_label, r.return_period): r.depth for r in estimate.results}
    for i, dur in enumerate(durations, start=1):
        row = table.rows[i].cells
        row[0].text = dur
        for j, rt in enumerate(return_periods, start=1):
            depth = lookup.get((dur, rt))
            row[j].text = f"{depth:.1f}" if depth is not None else "-"

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
